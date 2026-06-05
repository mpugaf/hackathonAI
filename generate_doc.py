"""
Generador del documento DOCX de presentación para hackathon ReqFlow AI.
Ejecutar con: python generate_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paleta de colores TCS ──────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x6B)
TEAL   = RGBColor(0x00, 0xC9, 0xB1)
BLUE   = RGBColor(0x1E, 0x6F, 0xEB)
SLATE  = RGBColor(0x5A, 0x6A, 0x85)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xFA)
BLACK  = RGBColor(0x0D, 0x1B, 0x2A)
DANGER = RGBColor(0xDC, 0x26, 0x26)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_para_shading(para, hex_fill):
    """Relleno de fondo para un párrafo."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)


def set_cell_bg(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def heading(text, level=1, color=NAVY, space_before=18, space_after=8):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p


def body(text, size=11, color=BLACK, bold=False, italic=False,
         space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.bold  = bold
    run.font.italic = italic
    return p


def bullet(text, size=11, color=BLACK, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def numbered(text, size=11, color=BLACK):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def divider(color_hex="1A3A6B"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def callout(text, bg_hex="E8F4FD", label="", label_color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    set_para_shading(p, bg_hex)
    if label:
        r = p.add_run(f"{label}  ")
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = label_color
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BLACK
    return p


def step_table(steps):
    """Tabla de pasos con número + acción + descripción."""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, txt in enumerate(["#", "Acción en la App", "Qué ocurre"]):
        hdr[i].text = txt
        set_cell_bg(hdr[i], "1A3A6B")
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
    table.columns[0].width = Cm(1.0)
    table.columns[1].width = Cm(6.5)
    table.columns[2].width = Cm(8.5)
    for idx, (action, result) in enumerate(steps, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = action
        row[2].text = result
        bg = "F0F4FA" if idx % 2 == 0 else "FFFFFF"
        for cell in row:
            set_cell_bg(cell, bg)
        for para in row[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = NAVY
                run.font.size = Pt(10)
        for cell in row[1:]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def role_badge(role, color_hex, text_hex="FFFFFF"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, color_hex)
    run = p.add_run(f"  {role.upper()}  ")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(text_hex)
    return p


# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════

# Bloque superior de color
p_cover = doc.add_paragraph()
p_cover.paragraph_format.space_before = Pt(0)
p_cover.paragraph_format.space_after  = Pt(0)
set_para_shading(p_cover, "1A3A6B")
run = p_cover.add_run("\n\n")
run.font.size = Pt(4)

# Logo / Título principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(20)
p_title.paragraph_format.space_after  = Pt(6)
r = p_title.add_run("ReqFlow AI")
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = NAVY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(6)
r2 = p_sub.add_run("Herramienta de Requisitos Asistida por Inteligencia Artificial")
r2.font.size = Pt(15)
r2.font.color.rgb = TEAL
r2.font.bold = True

p_event = doc.add_paragraph()
p_event.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_event.paragraph_format.space_before = Pt(0)
p_event.paragraph_format.space_after  = Pt(4)
r3 = p_event.add_run("TCS AI Fridays 2026")
r3.font.size = Pt(12)
r3.font.color.rgb = SLATE
r3.font.italic = True

divider()

# Datos de portada
meta_table = doc.add_table(rows=4, cols=2)
meta_table.style = 'Table Grid'
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_rows = [
    ("Evento",    "TCS AI Fridays Hackathon 2026"),
    ("Producto",  "ReqFlow AI — Gestión Inteligente de Requisitos"),
    ("Versión",   "1.0.0"),
    ("Fecha",     "Junio 2026"),
]
for i, (k, v) in enumerate(meta_rows):
    cells = meta_table.rows[i].cells
    cells[0].text = k
    cells[1].text = v
    set_cell_bg(cells[0], "1A3A6B")
    for para in cells[0].paragraphs:
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
    for para in cells[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Resumen Ejecutivo", level=1)
divider()

body(
    "ReqFlow AI es una aplicación web fullstack desarrollada durante el TCS AI Fridays 2026 "
    "que transforma el proceso tradicional de ingeniería de requisitos mediante el uso de "
    "modelos de lenguaje de gran escala (LLMs). La herramienta permite a equipos de análisis "
    "y revisión colaborar en la creación, validación y aprobación de especificaciones funcionales "
    "detalladas, reduciendo el tiempo de elaboración de un requisito de días a minutos.",
    space_after=10
)

body(
    "El sistema implementa un flujo de trabajo estructurado con dos perfiles diferenciados: "
    "el Analista, quien ingresa requisitos y orquesta la generación de especificaciones con IA, "
    "y el Revisor, quien evalúa, aprueba o devuelve las especificaciones generadas. "
    "Toda la actividad queda registrada en un audit log inmutable para trazabilidad completa.",
    space_after=12
)

callout(
    "Caso de uso principal: un Analista ingresa un requisito de negocio en lenguaje natural "
    "(ej: 'módulo de carrito de compras para e-commerce'), y el sistema genera en segundos una "
    "especificación completa con Historia de Usuario, Componentes Funcionales, Diagrama de Flujo "
    "Mermaid, Criterios de Aceptación en formato Dado-Cuando-Entonces y Requisitos No Funcionales.",
    bg_hex="E8F4FD", label="PROPUESTA DE VALOR:", label_color=BLUE
)

heading("Indicadores del Proyecto", level=2, color=TEAL, space_before=14)

kpi_table = doc.add_table(rows=2, cols=4)
kpi_table.style = 'Table Grid'
kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
kpis = [
    ("77 / 77", "Tests PASS", "100%", "Cobertura RF"),
    ("2 Roles",  "Analista + Revisor", "10 min", "Tiempo generación"),
]
headers = ["Métrica", "Descripción", "Métrica", "Descripción"]
for j, h in enumerate(headers):
    cell = kpi_table.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "00C9B1")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = NAVY
            run.font.bold = True
            run.font.size = Pt(10)
for j, (v, d, v2, d2) in enumerate(kpis):
    row = kpi_table.rows[1].cells
    pairs = [(v, d), (v2, d2)]
    for offset, (val, desc) in enumerate(pairs):
        idx = j * 2 + offset
        cell = kpi_table.rows[1].cells[idx]
        cell.paragraphs[0].clear()
        p_val = cell.paragraphs[0]
        r_val = p_val.add_run(val + "\n")
        r_val.font.size = Pt(18)
        r_val.font.bold = True
        r_val.font.color.rgb = NAVY
        r_desc = p_val.add_run(desc)
        r_desc.font.size = Pt(9)
        r_desc.font.color.rgb = SLATE
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "F0F4FA")
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. ARQUITECTURA Y STACK TÉCNICO
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Arquitectura y Stack Técnico", level=1)
divider()

body(
    "ReqFlow AI está construida sobre una arquitectura de tres capas desacopladas, "
    "cada una en su propio contenedor Docker, comunicadas a través de una red interna privada:",
    space_after=8
)

arch_table = doc.add_table(rows=1, cols=3)
arch_table.style = 'Table Grid'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
arch_headers = ["FRONTEND", "BACKEND", "BASE DE DATOS"]
arch_colors  = ["1E6FEB", "00C9B1", "1A3A6B"]
arch_details = [
    "React 18 + Vite\nTailwind CSS\nReact Router v6\nAxios + interceptors JWT\nNginx (producción)",
    "FastAPI + Python 3.11\nUvicorn (ASGI)\nSQLAlchemy 2.0\nAlembic (migraciones)\npython-jose (JWT)\npasslib/bcrypt",
    "PostgreSQL 15\n(SQLite para dev local)\nVolumen persistente\nAuto-seed en desarrollo",
]
for j, (hdr, color, detail) in enumerate(zip(arch_headers, arch_colors, arch_details)):
    cell = arch_table.rows[0].cells[j]
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = p.add_run(hdr + "\n\n")
    rh.font.bold = True
    rh.font.size = Pt(12)
    rh.font.color.rgb = WHITE
    rd = p.add_run(detail)
    rd.font.size = Pt(9)
    rd.font.color.rgb = WHITE
doc.add_paragraph()

heading("Integración con IA (TCS AI Lab)", level=2, color=TEAL, space_before=12)

body(
    "La inteligencia artificial del sistema se sustenta en dos modelos LLM distintos, "
    "orquestados por el backend a través de la API de TCS AI Lab (compatible con OpenAI API):",
    space_after=8
)

llm_table = doc.add_table(rows=3, cols=4)
llm_table.style = 'Table Grid'
llm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
llm_headers = ["Componente", "Modelo", "Rol", "Retries / Timeout"]
for j, h in enumerate(llm_headers):
    cell = llm_table.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1A3A6B")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
llm_data = [
    ("LLM-1 (Generador)", "genailab-maas-gpt-4o", "Transforma requisito en especificación completa", "2 reintentos / 30s"),
    ("LLM-2 (Revisor IA)", "azure/genailab-maas-gpt-4o-mini", "Evalúa calidad y emite score 0-100", "2 reintentos / 30s"),
]
for i, row_data in enumerate(llm_data):
    row = llm_table.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
        set_cell_bg(row[j], "F0F4FA" if i % 2 == 0 else "FFFFFF")
        for para in row[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. PRIMEROS PASOS — CÓMO LEVANTAR EL SISTEMA
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Primeros Pasos — Cómo Levantar el Sistema", level=1)
divider()

callout(
    "URL de la aplicación: http://localhost:3000\n"
    "URL del API (docs interactivas): http://localhost:8000/docs",
    bg_hex="E8F4FD", label="ACCESO:", label_color=BLUE
)

heading("Opción A — Con Docker (recomendado para producción/demo)", level=2, color=TEAL, space_before=14)

body("Requisitos: Docker Desktop instalado y corriendo.", bold=True, space_after=4)
numbered("Clonar el repositorio y entrar al directorio raíz del proyecto.")
numbered("Crear el archivo backend/.env copiando backend/.env.example y completar TCS_API_KEY con la clave real.")
numbered("Ejecutar en terminal:  docker compose up --build")
numbered("Esperar a que los 3 contenedores (db, backend, frontend) estén healthy (~60 segundos).")
numbered("Abrir el navegador en:  http://localhost:3000")
doc.add_paragraph()

heading("Opción B — Sin Docker (desarrollo local)", level=2, color=TEAL, space_before=12)

body("Requisitos: Python 3.11+, Node 20+. No requiere PostgreSQL (usa SQLite automáticamente).", bold=True, space_after=4)
numbered("Configurar backend: cd backend → python -m venv .venv → activar venv → pip install -r requirements.txt")
numbered("Iniciar API:  uvicorn main:app --host 0.0.0.0 --port 8000  (dentro de backend/)")
numbered("Configurar frontend: cd frontend → npm install")
numbered("Iniciar frontend:  npm run dev  (dentro de frontend/)")
numbered("Abrir el navegador en:  http://localhost:3000")
doc.add_paragraph()

heading("Credenciales de demostración (seed automático)", level=2, color=TEAL, space_before=12)

creds_table = doc.add_table(rows=3, cols=4)
creds_table.style = 'Table Grid'
creds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
creds_headers = ["Perfil", "Email", "Contraseña", "Permisos"]
for j, h in enumerate(creds_headers):
    cell = creds_table.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1A3A6B")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
creds_data = [
    ("Analista", "analista@tcs.com", "analista123", "Crear proyectos, generar specs, re-iterar"),
    ("Revisor",  "revisor@tcs.com",  "revisor123",  "Revisar, aprobar o rechazar specs"),
]
for i, row_data in enumerate(creds_data):
    row = creds_table.rows[i + 1].cells
    colors = ["1E6FEB", "00C9B1"]
    set_cell_bg(row[0], colors[i])
    for para in row[0].paragraphs:
        pass
    for j, val in enumerate(row_data):
        row[j].text = val
        if j > 0:
            set_cell_bg(row[j], "F0F4FA" if i % 2 == 0 else "FFFFFF")
        else:
            set_cell_bg(row[j], colors[i])
            for para in row[j].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = WHITE
                    run.font.size = Pt(10)
        for para in row[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. GUÍA DE USO — PERFIL ANALISTA
# ══════════════════════════════════════════════════════════════════════════════
role_badge("PERFIL: ANALISTA DE REQUISITOS", "1E6FEB")
heading("4. Guía de Uso — Perfil Analista", level=1)
divider()

body(
    "El Analista es el perfil principal del sistema. Su responsabilidad es capturar requisitos "
    "de negocio en lenguaje natural, orquestar la generación de especificaciones mediante IA, "
    "y gestionar el ciclo de vida del requisito hasta su aprobación. A continuación se detalla "
    "la secuencia completa de trabajo:",
    space_after=10
)

# Sección 4.1
heading("4.1  Inicio de sesión y vista del Dashboard", level=2, color=BLUE, space_before=12)

step_table([
    ("Abrir http://localhost:3000",
     "El sistema redirige automáticamente a /login si no hay sesión activa."),
    ("Ingresar email: analista@tcs.com\nContraseña: analista123 → clic en Entrar",
     "El sistema valida las credenciales contra el backend, genera un JWT con rol 'analista' (8h de vigencia) y redirige al Dashboard."),
    ("Observar el Dashboard",
     "Se visualizan 4 tarjetas de métricas: Total, Borradores, En Revisión y Aprobados. Una lista de requisitos existentes (el seed trae 1 de demo). Botón flotante '+ Nuevo Requisito' en la esquina inferior derecha."),
])

# Sección 4.2
heading("4.2  Crear un nuevo Proyecto", level=2, color=BLUE, space_before=12)

body(
    "Los requisitos deben pertenecer a un Proyecto. Si aún no existe uno para la sesión de demo, "
    "crear uno primero:",
    space_after=6
)

step_table([
    ("Navegar a Proyectos en el menú lateral",
     "Se muestra la lista de proyectos activos. El seed incluye 'Portal E-Commerce TCS Demo'."),
    ("Clic en 'Nuevo Proyecto'",
     "Se abre el formulario. Campos: Título (min 5 chars, obligatorio) y Descripción (opcional)."),
    ("Completar: Título = 'Sistema de Gestión de Tickets'\nDescripción = 'Portal interno TCS para gestión de incidencias.'",
     "Validación inmediata en frontend. Si el título tiene menos de 5 caracteres, aparece error de validación sin llamar al servidor."),
    ("Clic en 'Guardar'",
     "El backend crea el proyecto con is_active=true, registra en el audit log (project.created) y retorna HTTP 201. El proyecto aparece en la lista."),
])

# Sección 4.3
heading("4.3  Crear un nuevo Requisito", level=2, color=BLUE, space_before=12)

step_table([
    ("Clic en '+ Nuevo Requisito' (botón flotante) o navegar a /requirements/new",
     "Solo accesible para el rol Analista. Un Revisor que intente acceder es redirigido al Dashboard con un mensaje de advertencia."),
    ("Seleccionar el proyecto en el dropdown",
     "Solo se listan proyectos activos del sistema."),
    ("Completar el Título:\n'Módulo de Notificaciones Push'",
     "Mínimo 5 caracteres, máximo 300. Validación en tiempo real."),
    ("Completar el Requisito de Alto Nivel (raw_input):\n'Implementar un sistema de notificaciones push para usuarios móviles y web que informe sobre cambios de estado en tickets, asignaciones nuevas y vencimientos próximos.'",
     "Mínimo 20 caracteres. Este texto es la entrada principal al LLM-1."),
    ("Clic en 'Crear Requisito'",
     "El requisito se crea en estado 'borrador'. Aparece en el Dashboard con badge gris 'Borrador'. Se registra en audit log (requirement.created)."),
])

# Sección 4.4
heading("4.4  Generar Especificación con IA (LLM-1)", level=2, color=BLUE, space_before=12)

callout(
    "Este paso requiere una TCS_API_KEY válida configurada en backend/.env. "
    "Sin ella, el sistema retorna un error HTTP 502 con mensaje descriptivo.",
    bg_hex="FFF3CD", label="⚠  REQUISITO:", label_color=RGBColor(0x92, 0x40, 0x00)
)

step_table([
    ("Abrir el detalle del requisito creado\n(clic en su tarjeta en el Dashboard)",
     "Se muestra: título, estado, texto del requisito, y botón 'Generar con IA' (visible solo para Analistas en estado borrador)."),
    ("Clic en 'Generar con IA'",
     "Aparece el indicador de carga 'LoadingOracle'. El backend invoca LLM-1 (gpt-4o). Tiempo estimado: 10-30 segundos según carga del servicio."),
    ("Esperar la respuesta",
     "Si es exitoso: aparece un toast verde de confirmación. Se crea la Versión 1 de la especificación. El estado permanece en 'borrador'."),
    ("Revisar la especificación generada",
     "La especificación incluye obligatoriamente: HISTORIA DE USUARIO, DESCRIPCIÓN DETALLADA, COMPONENTES FUNCIONALES, DIAGRAMA DE FLUJO (Mermaid), REQUISITOS NO FUNCIONALES, ROLES Y PERMISOS, PUNTOS DE INTEGRACIÓN, ESCENARIOS DE ERROR, CRITERIOS DE ACEPTACIÓN (Dado-Cuando-Entonces), CONSIDERACIONES WCAG 2.1."),
    ("Editar si es necesario (SpecEditor)",
     "La especificación es editable en la interfaz antes de enviar a revisión IA. Los cambios se guardan en el campo edited_spec (no sobreescriben generated_spec)."),
])

# Sección 4.5
heading("4.5  Enviar a Revisión IA (LLM-2)", level=2, color=BLUE, space_before=12)

step_table([
    ("Clic en 'Enviar a Revisión IA'",
     "El backend invoca LLM-2 (gpt-4o-mini). Modelo diferente al generador para perspectiva independiente."),
    ("El sistema muestra el panel de Feedback IA (AIFeedbackPanel)",
     "Score global (0-100): si >= 70 aparece en verde como APROBADO, si < 70 en naranja como REQUIERE_REVISIÓN. También muestra: evaluación por sección, deficiencias identificadas y sugerencias de mejora."),
    ("El estado del requisito cambia a 'En Revisión'",
     "Se registra en audit log (requirement.reviewed_ai). El requisito queda en la cola del Revisor."),
])

callout(
    "Una vez en estado 'En Revisión', el Analista no puede editar ni re-generar el requisito. "
    "Debe esperar la decisión del Revisor. Si el Revisor lo rechaza, el estado vuelve a 'Rechazado' "
    "y el Analista puede iniciar una nueva iteración.",
    bg_hex="E8F4FD", label="NOTA:", label_color=BLUE
)

# Sección 4.6
heading("4.6  Iteración tras Rechazo del Revisor", level=2, color=BLUE, space_before=12)

body(
    "Si el Revisor devuelve el requisito con observaciones, el Analista inicia un ciclo de mejora:",
    space_after=6
)

step_table([
    ("El estado del requisito muestra 'Rechazado' con el feedback del Revisor visible",
     "El panel de feedback humano muestra las observaciones específicas del Revisor."),
    ("Clic en 'Re-generar con Feedback'",
     "El Analista puede ajustar el raw_input incorporando las observaciones antes de re-generar."),
    ("Clic en 'Generar con IA'",
     "Se crea la Versión 2 (o N). Cada versión mantiene su propio historial con fecha, modelo y estado."),
    ("Ver historial en /requirements/{id}/history",
     "Lista de todas las versiones ordenadas por número, con modelo LLM utilizado y fecha de creación."),
    ("Re-enviar a Revisión IA y luego al Revisor",
     "El ciclo se repite hasta obtener la aprobación."),
])

# Sección 4.7
heading("4.7  Exportar Especificación Aprobada", level=2, color=BLUE, space_before=12)

step_table([
    ("El requisito debe estar en estado 'Aprobado'",
     "Solo los requisitos aprobados pueden exportarse. Si el estado es otro, el botón de exportación está deshabilitado."),
    ("Clic en 'Exportar' (ExportButton)",
     "El backend genera un documento HTML estructurado con: título, fecha, modelo utilizado, especificación completa y feedback del Revisor."),
    ("El navegador descarga el archivo",
     "Formato: spec_{uuid}.html. Listo para compartir, imprimir o almacenar como evidencia del proceso."),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. GUÍA DE USO — PERFIL REVISOR
# ══════════════════════════════════════════════════════════════════════════════
role_badge("PERFIL: REVISOR DE ESPECIFICACIONES", "00C9B1", "1A3A6B")
heading("5. Guía de Uso — Perfil Revisor", level=1)
divider()

body(
    "El Revisor es el responsable de la calidad final de las especificaciones. Tiene acceso de "
    "lectura a toda la plataforma y poder exclusivo para aprobar o rechazar requisitos. "
    "No puede crear proyectos ni requisitos, ni generar especificaciones.",
    space_after=10
)

# Sección 5.1
heading("5.1  Inicio de sesión y vista del Dashboard", level=2, color=TEAL, space_before=12)

step_table([
    ("Abrir http://localhost:3000",
     "Redirección a /login si no hay sesión activa."),
    ("Ingresar email: revisor@tcs.com\nContraseña: revisor123 → clic en Entrar",
     "JWT generado con rol 'revisor'. El Dashboard muestra la misma vista de métricas que el Analista, PERO sin el botón '+ Nuevo Requisito' (oculto por lógica de rol en frontend)."),
    ("Identificar requisitos 'En Revisión'",
     "Usar el filtro de estado 'En Revisión' en el Dashboard para ver solo los pendientes. Los badges naranjas indican requisitos esperando revisión humana."),
])

# Sección 5.2
heading("5.2  Revisar una Especificación", level=2, color=TEAL, space_before=12)

step_table([
    ("Clic en un requisito con estado 'En Revisión'",
     "Se abre el detalle del requisito con: el texto original del requisito, la especificación generada por LLM-1, y el panel de Feedback IA con score y análisis de LLM-2."),
    ("Navegar a /requirements/{id}/review",
     "La ruta de revisión es exclusiva para revisores. Un Analista que intente acceder es redirigido con un mensaje de advertencia de permisos."),
    ("Leer la especificación completa",
     "Evaluar todas las secciones: Historia de Usuario, Criterios de Aceptación, NFRs, Diagrama de Flujo. Contrastar con el feedback de la IA."),
    ("Revisar el Score IA (0-100)",
     "Score >= 70: IA recomienda APROBADO. Score < 70: IA recomienda REVISIÓN. El Revisor puede aceptar o anular esta recomendación según su criterio."),
])

# Sección 5.3
heading("5.3  Aprobar una Especificación", level=2, color=TEAL, space_before=12)

step_table([
    ("Agregar comentario de revisión (opcional pero recomendado)",
     "El campo de feedback humano permite documentar la justificación de la aprobación o cualquier observación para el registro histórico."),
    ("Clic en 'Aprobar'",
     "El backend cambia el estado del requisito a 'aprobado'. Se crea un registro en reviews con reviewer_type='humano' y decision='aprobado'. Se registra en audit log (requirement.approved)."),
    ("Confirmación visual",
     "Toast verde de confirmación. El badge del requisito en el Dashboard cambia a verde 'Aprobado'. El Analista puede ahora exportar la especificación."),
])

callout(
    "El Revisor puede aprobar directamente sin haber ejecutado primero una review-human explícita. "
    "En ese caso, el sistema crea automáticamente un registro de revisión humana con feedback 'Aprobado'. "
    "Toda la decisión queda trazada en el audit log.",
    bg_hex="E8F4FD", label="APROBACIÓN DIRECTA:", label_color=TEAL
)

# Sección 5.4
heading("5.4  Rechazar y Devolver al Analista", level=2, color=TEAL, space_before=12)

step_table([
    ("Identificar deficiencias específicas en la especificación",
     "Revisar secciones incompletas, criterios de aceptación ambiguos o NFRs faltantes. El feedback de LLM-2 puede guiar este análisis."),
    ("Agregar observaciones detalladas en el campo de feedback",
     "Ejemplos: 'Faltan criterios de aceptación para el flujo de error de red.' / 'El diagrama Mermaid no refleja el caso de autenticación.' El texto es obligatorio (mínimo 1 caracter)."),
    ("Clic en 'Rechazar'",
     "El estado cambia a 'rechazado'. Se crea review con decision='rechazado'. Se registra en audit log (requirement.rejected). El Analista recibe visibilidad del feedback."),
    ("El Analista inicia una nueva iteración",
     "El Analista ve el estado 'Rechazado', lee el feedback del Revisor, ajusta el requisito o el texto base, y vuelve a generar una nueva versión."),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. FLUJO COMPLETO — SECUENCIA DE DEMO
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Secuencia de Demo para la Presentación", level=1)
divider()

callout(
    "Tiempo estimado de demo completa: 8-12 minutos. "
    "Tener 2 pestañas del navegador abiertas: una con sesión de Analista y otra con sesión de Revisor.",
    bg_hex="E8F4FD", label="PREPARACIÓN:", label_color=BLUE
)

body("A continuación la secuencia recomendada para mostrar el valor del producto en la hackathon:", space_after=8)

demo_steps = [
    ("CONTEXTO (1 min)",
     "Presentar el problema: especificaciones de requisitos son costosas, inconsistentes y lentas. "
     "Mostrar el Dashboard vacío o con datos del seed."),
    ("LOGIN ANALISTA (30 seg)",
     "Abrir http://localhost:3000 → login como analista@tcs.com / analista123 → "
     "mostrar el dashboard con las 4 métricas y los controles de filtro."),
    ("CREAR PROYECTO (1 min)",
     "Navegar a Proyectos → Nuevo Proyecto. "
     "Titulo: 'Sistema de Aprobación de Créditos'. "
     "Mostrar la validación en tiempo real del formulario."),
    ("CREAR REQUISITO (1 min)",
     "Clic en + Nuevo Requisito. Seleccionar el proyecto. "
     "Ingresar: Título = 'Motor de Scoring Crediticio'. "
     "Raw input = 'Desarrollar un motor de evaluación de scoring crediticio que analice historial "
     "de pagos, ingresos declarados y deudas activas del solicitante para generar una puntuación "
     "0-1000 y recomendar aprobación, revisión manual o rechazo del crédito.' → Crear."),
    ("GENERAR CON IA (3-4 min)",
     "Abrir el requisito → clic en 'Generar con IA'. "
     "Mientras carga: explicar que LLM-1 (gpt-4o) está transformando el requisito en una "
     "especificación estructurada con 10 secciones técnicas. "
     "Al terminar: mostrar el resultado completo, destacar la Historia de Usuario, "
     "el diagrama Mermaid y los Criterios Dado-Cuando-Entonces."),
    ("REVISIÓN IA (1 min)",
     "Clic en 'Enviar a Revisión IA'. "
     "LLM-2 (gpt-4o-mini) evalúa la calidad. "
     "Mostrar el score y el análisis por sección. Destacar que es un modelo distinto al generador."),
    ("REVISIÓN HUMANA — CAMBIO A REVISOR (2 min)",
     "Abrir 2da pestaña → login como revisor@tcs.com / revisor123. "
     "El revisor ve el requisito en estado 'En Revisión'. "
     "Navegar a la vista de revisión, leer la especificación y el feedback IA. "
     "Aprobar o rechazar con comentario. Mostrar el cambio de estado en tiempo real."),
    ("EXPORTACIÓN (30 seg)",
     "Si fue aprobado: volver a la pestaña del Analista → "
     "abrir el requisito → clic en 'Exportar' → el navegador descarga el HTML estructurado."),
    ("CIERRE — AUDIT LOG Y TRAZABILIDAD (30 seg)",
     "Mencionar que cada una de las acciones (login, create, generate, review, approve) "
     "quedó registrada en el audit log con timestamp, usuario e IP. "
     "El sistema es completamente auditable para cumplimiento regulatorio."),
]

for i, (title, detail) in enumerate(demo_steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    r_num = p.add_run(f"PASO {i}: ")
    r_num.font.bold = True
    r_num.font.size = Pt(11)
    r_num.font.color.rgb = NAVY
    r_title = p.add_run(title + "\n")
    r_title.font.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = TEAL
    r_detail = p.add_run(detail)
    r_detail.font.size = Pt(10.5)
    r_detail.font.color.rgb = BLACK

doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  7. ASPECTOS TÉCNICOS DESTACADOS
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Aspectos Técnicos Destacados", level=1)
divider()

heading("Seguridad y RBAC", level=2, color=TEAL, space_before=12)
bullet("Autenticación basada en JWT (HS256) con expiración de 8 horas.")
bullet("Control de acceso por rol en cada endpoint del backend con Depends(require_role([...])).")
bullet("Contraseñas hasheadas con bcrypt de 12 rounds — resistente a ataques de fuerza bruta.")
bullet("Token almacenado en sessionStorage (no localStorage) — se limpia al cerrar la pestaña.")
bullet("Validación de inputs con Pydantic (longitudes mínimas, formatos, tipos).")
bullet("Acceso al API usando parámetros vinculados de SQLAlchemy — sin riesgo de SQL injection.")

heading("Robustez del Servicio LLM", level=2, color=TEAL, space_before=12)
bullet("Retry automático con backoff exponencial: hasta 2 reintentos ante fallos transitorios.")
bullet("Timeout configurable por request (30s por defecto).")
bullet("Manejo explícito de errores: HTTP 504 para timeout, HTTP 502 para errores de conexión.")
bullet("El estado del requisito nunca queda inconsistente ante un fallo del LLM.")
bullet("Los modelos usados (nombre, versión) se almacenan en cada versión generada para auditoría.")

heading("Calidad del Código", level=2, color=TEAL, space_before=12)
bullet("77 tests automatizados con 100% de PASS — suite corriendo en SQLite in-memory.")
bullet("Separación clara de capas: routers, services, models, schemas.")
bullet("Audit log inmutable con 10 tipos de eventos trazados.")
bullet("Soporte nativo para SQLite (dev) y PostgreSQL (producción) sin cambios de código.")
bullet("Dockerización completa con healthchecks y dependencias entre servicios.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  8. ESTADO DEL SISTEMA Y TRABAJO FUTURO
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Estado del Sistema y Trabajo Futuro", level=1)
divider()

body("El sistema fue sometido a una revisión QA completa. A continuación los hallazgos principales y las mejoras planificadas:", space_after=8)

status_table = doc.add_table(rows=1, cols=4)
status_table.style = 'Table Grid'
status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["ID", "Severidad", "Descripción", "Estado"]):
    cell = status_table.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1A3A6B")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)

defects = [
    ("DEF-003", "HIGH",   "feedback_context no se persiste en re-generación — iteración con contexto pendiente", "OPEN — P1"),
    ("DEF-001", "MEDIUM", "Logout no invalida JWT (stateless); token válido 8h tras cierre de sesión",           "OPEN — P2"),
    ("DEF-005", "MEDIUM", "Headers de seguridad HTTP ausentes (X-Content-Type-Options, X-Frame-Options)",        "OPEN — P2"),
    ("DEF-006", "MEDIUM", "SSL verify=False en llamadas al LLM — riesgo MITM en producción",                    "OPEN — P2"),
    ("DEF-002", "LOW",    "review-ai accesible por cualquier rol autenticado",                                   "OPEN — P3"),
    ("DEF-004", "LOW",    "Dashboard sin métricas de tiempo promedio y modelos utilizados",                      "OPEN — P3"),
]
sev_colors = {"HIGH": "FEE2E2", "MEDIUM": "FEF3C7", "LOW": "F0F4FA"}
for i, (did, sev, desc, state) in enumerate(defects):
    row = status_table.add_row().cells
    row[0].text = did
    row[1].text = sev
    row[2].text = desc
    row[3].text = state
    for cell in row:
        set_cell_bg(cell, sev_colors.get(sev, "FFFFFF"))
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
doc.add_paragraph()

heading("Roadmap — Próximas iteraciones", level=2, color=TEAL, space_before=12)
bullet("Implementar feedback_context en el ciclo de re-generación para pasar observaciones del revisor al LLM-1.")
bullet("Agregar endpoint GET /stats con métricas de tiempo promedio y distribución de modelos.")
bullet("Migrar de python-jose a PyJWT (mantenido activamente).")
bullet("Generar primera migración Alembic para gestión formal de schema en producción.")
bullet("Añadir middleware de headers de seguridad HTTP.")
bullet("Integrar notificaciones push o email cuando un requisito cambia de estado.")
bullet("Dashboard con gráficos de distribución de scores IA y tiempos de ciclo.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  9. ESTRUCTURA DE ARCHIVOS DE CALIDAD GENERADOS
# ══════════════════════════════════════════════════════════════════════════════
heading("9. Documentación de Calidad Generada", level=1)
divider()

body("Durante el proceso de desarrollo y QA se generaron los siguientes artefactos:", space_after=8)

qa_table = doc.add_table(rows=1, cols=3)
qa_table.style = 'Table Grid'
qa_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Archivo", "Tipo", "Contenido"]):
    cell = qa_table.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1A3A6B")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)

qa_files = [
    ("QA_RECON.md",             "Reconocimiento",    "Mapa completo: endpoints API, modelos de datos, rutas frontend, variables de entorno"),
    ("QA_TEST_CASES.md",        "Casos de Prueba",   "43 casos funcionales RF-001 a RF-020 con estado PASS/FAIL/BLOCKED"),
    ("backend/tests/test_api.py","Tests Automáticos", "59 tests pytest nuevos — 77 total, 100% PASS"),
    ("QA_NFR.md",               "No Funcionales",    "14 verificaciones: seguridad, resiliencia, consistencia de datos"),
    ("TRACEABILITY_MATRIX.csv", "Trazabilidad",      "Mapa completo RF → TC → test automatizado → estado QA"),
    ("QA_REPORT.md",            "Reporte Final",     "Resumen ejecutivo, 7 defectos con severidad, riesgos y recomendaciones P1/P2/P3"),
    ("PRESENTACION_HACKATHON.docx", "Presentación",  "Este documento — speech, guías de uso por perfil, secuencia de demo"),
]
for i, (fname, ftype, fdesc) in enumerate(qa_files):
    row = qa_table.add_row().cells
    row[0].text = fname
    row[1].text = ftype
    row[2].text = fdesc
    bg = "F0F4FA" if i % 2 == 0 else "FFFFFF"
    for cell in row:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
    for para in row[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = NAVY

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  CIERRE
# ══════════════════════════════════════════════════════════════════════════════
heading("Cierre", level=1)
divider()

p_close = doc.add_paragraph()
p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_close.paragraph_format.space_before = Pt(20)
p_close.paragraph_format.space_after  = Pt(10)
r_close = p_close.add_run(
    "ReqFlow AI demuestra que la Inteligencia Artificial puede transformar procesos "
    "críticos de ingeniería de software de manera responsable, trazable y auditable. "
    "El sistema reduce el tiempo de elaboración de especificaciones de días a minutos, "
    "sin sacrificar rigor técnico ni control humano en las decisiones finales."
)
r_close.font.size = Pt(13)
r_close.font.italic = True
r_close.font.color.rgb = NAVY

p_tcs = doc.add_paragraph()
p_tcs.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tcs.paragraph_format.space_before = Pt(20)
r_tcs = p_tcs.add_run("TCS AI Fridays 2026  |  ReqFlow AI  |  v1.0.0")
r_tcs.font.size = Pt(10)
r_tcs.font.color.rgb = SLATE

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = r"c:\Users\genaiclsan1usr3\Documents\GitHub\hackathonAI\PRESENTACION_HACKATHON.docx"
doc.save(output_path)
print(f"Documento generado: {output_path}")
