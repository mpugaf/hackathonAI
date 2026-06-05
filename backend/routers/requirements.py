import html
import uuid

import requests
from fastapi import APIRouter, Depends, HTTPException, Request, Response, status
from sqlalchemy import func
from sqlalchemy.orm import Session

from db.database import get_db
from models.project import Project
from models.requirement import Requirement
from models.requirement_version import RequirementVersion
from models.review import Review
from models.user import User
from schemas.requirement import RequirementCreate, RequirementDetail, RequirementOut, RequirementUpdate
from schemas.requirement_version import VersionOut
from schemas.review import RejectRequest, ReviewHumanRequest, ReviewOut
from services import llm_service
from services.audit_service import log_action
from services.auth_service import get_current_user, require_role

router = APIRouter()


def _client_ip(request: Request) -> str | None:
    return request.client.host if request.client else None


def _get_requirement(db: Session, requirement_id: uuid.UUID) -> Requirement:
    requirement = db.get(Requirement, requirement_id)
    if not requirement:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Requisito no encontrado")
    return requirement


def _latest_version(db: Session, requirement_id: uuid.UUID) -> RequirementVersion | None:
    return (
        db.query(RequirementVersion)
        .filter(RequirementVersion.requirement_id == requirement_id)
        .order_by(RequirementVersion.version_number.desc())
        .first()
    )


def _latest_review(db: Session, version_id: uuid.UUID, reviewer_type: str) -> Review | None:
    return (
        db.query(Review)
        .filter(Review.requirement_version_id == version_id, Review.reviewer_type == reviewer_type)
        .order_by(Review.created_at.desc())
        .first()
    )


def _to_requirement_out(db: Session, requirement: Requirement) -> RequirementOut:
    version_count = db.query(func.count(RequirementVersion.id)).filter(RequirementVersion.requirement_id == requirement.id).scalar() or 0
    return RequirementOut(
        id=requirement.id,
        project_id=requirement.project_id,
        title=requirement.title,
        raw_input=requirement.raw_input,
        status=requirement.status,
        created_by=requirement.created_by,
        created_at=requirement.created_at,
        updated_at=requirement.updated_at,
        version_count=version_count,
    )


@router.get("", response_model=list[RequirementOut])
def list_requirements(
    project_id: uuid.UUID | None = None,
    status: str | None = None,
    skip: int = 0,
    limit: int = 20,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> list[RequirementOut]:
    del current_user
    query = db.query(Requirement)
    if project_id:
        query = query.filter(Requirement.project_id == project_id)
    if status:
        query = query.filter(Requirement.status == status)
    requirements = query.order_by(Requirement.created_at.desc()).offset(skip).limit(limit).all()
    return [_to_requirement_out(db, item) for item in requirements]


@router.post("", response_model=RequirementOut, status_code=status.HTTP_201_CREATED)
def create_requirement(
    payload: RequirementCreate,
    request: Request,
    current_user: User = Depends(require_role(["analista"])),
    db: Session = Depends(get_db),
) -> RequirementOut:
    project = db.get(Project, payload.project_id)
    if not project or not project.is_active:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="project_id no existe o esta inactivo")
    requirement = Requirement(
        project_id=payload.project_id,
        title=payload.title,
        raw_input=payload.raw_input,
        created_by=current_user.id,
        status="borrador",
    )
    db.add(requirement)
    db.commit()
    db.refresh(requirement)
    log_action(db, current_user.id, "requirement.created", "requirement", requirement.id, ip_address=_client_ip(request))
    return _to_requirement_out(db, requirement)


@router.get("/{requirement_id}", response_model=RequirementDetail)
def get_requirement_detail(
    requirement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> RequirementDetail:
    del current_user
    requirement = _get_requirement(db, requirement_id)
    latest = _latest_version(db, requirement.id)
    latest_ai_review = _latest_review(db, latest.id, "ia") if latest else None
    latest_human_review = _latest_review(db, latest.id, "humano") if latest else None
    return RequirementDetail(
        id=requirement.id,
        project_id=requirement.project_id,
        title=requirement.title,
        raw_input=requirement.raw_input,
        status=requirement.status,
        created_by=requirement.created_by,
        created_at=requirement.created_at,
        updated_at=requirement.updated_at,
        latest_version=latest,
        latest_ai_review=latest_ai_review,
        latest_human_review=latest_human_review,
    )


@router.put("/{requirement_id}", response_model=RequirementOut)
def update_requirement(
    requirement_id: uuid.UUID,
    payload: RequirementUpdate,
    request: Request,
    current_user: User = Depends(require_role(["analista"])),
    db: Session = Depends(get_db),
) -> RequirementOut:
    requirement = _get_requirement(db, requirement_id)
    if requirement.status != "borrador":
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail=f"No se puede editar un requisito en estado {requirement.status}")
    if payload.title is not None:
        requirement.title = payload.title
    if payload.raw_input is not None:
        requirement.raw_input = payload.raw_input
    db.commit()
    db.refresh(requirement)
    log_action(db, current_user.id, "requirement.updated", "requirement", requirement.id, ip_address=_client_ip(request))
    return _to_requirement_out(db, requirement)


@router.post("/{requirement_id}/generate")
def generate_requirement(
    requirement_id: uuid.UUID,
    request: Request,
    current_user: User = Depends(require_role(["analista"])),
    db: Session = Depends(get_db),
) -> dict:
    requirement = _get_requirement(db, requirement_id)
    max_version = db.query(func.max(RequirementVersion.version_number)).filter(RequirementVersion.requirement_id == requirement.id).scalar()
    version_number = (max_version or 0) + 1
    try:
        result = llm_service.generate_spec(requirement.raw_input)
    except requests.exceptions.Timeout as exc:
        raise HTTPException(status_code=status.HTTP_504_GATEWAY_TIMEOUT, detail="Timeout al llamar al LLM") from exc
    except requests.exceptions.RequestException as exc:
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail=f"Error al llamar al LLM: {exc}") from exc

    version = RequirementVersion(
        requirement_id=requirement.id,
        version_number=version_number,
        generated_spec=result["content"],
        model_used=result["model"],
        prompt_used=result["prompt_used"],
    )
    requirement.status = "borrador"
    db.add(version)
    db.commit()
    db.refresh(version)
    log_action(
        db,
        current_user.id,
        "requirement.generated",
        "requirement",
        requirement.id,
        metadata={"model": result["model"], "version_number": version_number, "duration_ms": result["duration_ms"]},
        ip_address=_client_ip(request),
    )
    return {"version_number": version.version_number, "generated_spec": version.generated_spec, "model_used": version.model_used, "created_at": version.created_at}


@router.post("/{requirement_id}/review-ai")
def review_requirement_ai(
    requirement_id: uuid.UUID,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    requirement = _get_requirement(db, requirement_id)
    version = _latest_version(db, requirement.id)
    if not version:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Primero genera la especificacion")
    try:
        result = llm_service.review_spec(requirement.raw_input, version.generated_spec)
    except requests.exceptions.Timeout as exc:
        raise HTTPException(status_code=status.HTTP_504_GATEWAY_TIMEOUT, detail="Timeout al llamar al LLM") from exc
    except requests.exceptions.RequestException as exc:
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail=f"Error al llamar al LLM: {exc}") from exc

    review = Review(
        requirement_version_id=version.id,
        reviewer_id=None,
        reviewer_type="ia",
        feedback=result["content"],
        ai_score=result["ai_score"],
        decision="pendiente",
    )
    requirement.status = "en_revision"
    db.add(review)
    db.commit()
    log_action(db, current_user.id, "requirement.reviewed_ai", "requirement", requirement.id, ip_address=_client_ip(request))
    return {"feedback": review.feedback, "ai_score": review.ai_score}


@router.post("/{requirement_id}/review-human", response_model=ReviewOut)
def review_requirement_human(
    requirement_id: uuid.UUID,
    payload: ReviewHumanRequest,
    request: Request,
    current_user: User = Depends(require_role(["revisor"])),
    db: Session = Depends(get_db),
) -> Review:
    requirement = _get_requirement(db, requirement_id)
    version = _latest_version(db, requirement.id)
    if not version:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Primero genera la especificacion")
    if not _latest_review(db, version.id, "ia"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="La revision IA debe ejecutarse primero")
    review = Review(
        requirement_version_id=version.id,
        reviewer_id=current_user.id,
        reviewer_type="humano",
        feedback=payload.feedback,
        decision=payload.decision,
    )
    db.add(review)
    db.commit()
    db.refresh(review)
    log_action(db, current_user.id, "requirement.reviewed_human", "review", review.id, ip_address=_client_ip(request))
    return review


@router.post("/{requirement_id}/approve")
def approve_requirement(
    requirement_id: uuid.UUID,
    request: Request,
    current_user: User = Depends(require_role(["revisor"])),
    db: Session = Depends(get_db),
) -> dict:
    requirement = _get_requirement(db, requirement_id)
    if requirement.status != "en_revision":
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Solo se pueden aprobar requisitos en revision")
    version = _latest_version(db, requirement.id)
    if not version:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Primero genera la especificacion")
    review = _latest_review(db, version.id, "humano")
    if review:
        review.decision = "aprobado"
    else:
        db.add(Review(requirement_version_id=version.id, reviewer_id=current_user.id, reviewer_type="humano", feedback="Aprobado", decision="aprobado"))
    requirement.status = "aprobado"
    db.commit()
    log_action(db, current_user.id, "requirement.approved", "requirement", requirement.id, ip_address=_client_ip(request))
    return {"status": "aprobado", "message": "Requisito aprobado exitosamente"}


@router.post("/{requirement_id}/reject")
def reject_requirement(
    requirement_id: uuid.UUID,
    payload: RejectRequest,
    request: Request,
    current_user: User = Depends(require_role(["revisor"])),
    db: Session = Depends(get_db),
) -> dict:
    requirement = _get_requirement(db, requirement_id)
    if requirement.status != "en_revision":
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Solo se pueden rechazar requisitos en revision")
    version = _latest_version(db, requirement.id)
    if not version:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Primero genera la especificacion")
    db.add(Review(requirement_version_id=version.id, reviewer_id=current_user.id, reviewer_type="humano", feedback=payload.feedback, decision="rechazado"))
    requirement.status = "rechazado"
    db.commit()
    log_action(db, current_user.id, "requirement.rejected", "requirement", requirement.id, ip_address=_client_ip(request))
    return {"status": "rechazado", "feedback": payload.feedback}


@router.get("/{requirement_id}/export")
def export_requirement(
    requirement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> Response:
    del current_user
    requirement = _get_requirement(db, requirement_id)
    if requirement.status != "aprobado":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Solo se pueden exportar requisitos aprobados")
    version = _latest_version(db, requirement.id)
    if not version:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No existe version generada")
    human_review = _latest_review(db, version.id, "humano")
    body = f"""<!doctype html>
<html lang="es"><head><meta charset="utf-8"><title>{html.escape(requirement.title)}</title></head>
<body>
<h1>{html.escape(requirement.title)}</h1>
<p><strong>Fecha:</strong> {version.created_at}</p>
<p><strong>Modelo:</strong> {html.escape(version.model_used)}</p>
<pre>{html.escape(version.generated_spec)}</pre>
<p><strong>Revision:</strong> {html.escape(human_review.feedback if human_review else "Aprobado")}</p>
</body></html>"""
    return Response(
        content=body,
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="spec_{requirement.id}.html"'},
    )


@router.get("/{requirement_id}/export-docx")
def export_requirement_docx(
    requirement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> Response:
    del current_user
    import io
    import re as _re
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    requirement = _get_requirement(db, requirement_id)
    version = _latest_version(db, requirement.id)
    if not version:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No existe version generada")

    doc = DocxDocument()
    for section in doc.sections:
        from docx.shared import Cm
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Título principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(requirement.title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

    # Metadatos
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(f"Versión {version.version_number}  ·  Modelo: {version.model_used}  ·  {version.created_at.strftime('%d/%m/%Y %H:%M')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x5A, 0x6A, 0x85)
    meta_run.italic = True

    doc.add_paragraph()

    # Parsear el markdown de la spec y convertir a estilos DOCX
    spec_text = version.generated_spec or ""
    lines = spec_text.splitlines()

    def strip_inline(text):
        """Devuelve lista de (texto, bold, italic) para una línea."""
        segments = []
        pattern = _re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|__(.+?)__|_(.+?)_|\*(.+?)\*|`(.+?)`)')
        last = 0
        for m in pattern.finditer(text):
            if m.start() > last:
                segments.append((text[last:m.start()], False, False))
            raw = m.group(0)
            inner = m.group(2) or m.group(3) or m.group(4) or m.group(5) or m.group(6) or m.group(7)
            bold   = '**' in raw or '__' in raw
            italic = raw.startswith('*') and not raw.startswith('**') or raw.startswith('_') and not raw.startswith('__')
            mono   = raw.startswith('`')
            segments.append((inner, bold or ('***' in raw), italic and not bold, mono))
            last = m.end()
        if last < len(text):
            segments.append((text[last:], False, False))
        return segments

    def add_styled_paragraph(p, line):
        for seg in strip_inline(line):
            if len(seg) == 3:
                txt, bold, italic = seg
                mono = False
            else:
                txt, bold, italic, mono = seg
            run = p.add_run(txt)
            run.bold = bold
            run.italic = italic
            if mono:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x6F, 0xEB)

    for line in lines:
        stripped = line.rstrip()

        # H1
        if _re.match(r'^#\s+', stripped):
            text = _re.sub(r'^#+\s*', '', stripped)
            p = doc.add_heading(level=1)
            p.clear()
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
            continue

        # H2
        if _re.match(r'^#{2}\s+', stripped):
            text = _re.sub(r'^#+\s*', '', stripped)
            p = doc.add_heading(level=2)
            p.clear()
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x6F, 0xEB)
            continue

        # H3
        if _re.match(r'^#{3,}\s+', stripped):
            text = _re.sub(r'^#+\s*', '', stripped)
            p = doc.add_heading(level=3)
            p.clear()
            run = p.add_run(text)
            run.bold = True
            continue

        # Secciones numeradas sin #: "1. TITULO EN MAYÚSCULAS"
        if _re.match(r'^\d{1,2}\.\s+[A-ZÁÉÍÓÚÜÑ\s]{4,}$', stripped):
            p = doc.add_heading(level=2)
            p.clear()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
            continue

        # Bullet -/*/+
        if _re.match(r'^[\-\*\+]\s+', stripped):
            text = _re.sub(r'^[\-\*\+]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            add_styled_paragraph(p, text)
            continue

        # Bullet numerado
        if _re.match(r'^\d+\.\s+', stripped) and not _re.match(r'^\d+\.\s+[A-ZÁÉÍÓÚ\s]{4,}$', stripped):
            text = _re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            add_styled_paragraph(p, text)
            continue

        # Línea en blanco
        if not stripped:
            doc.add_paragraph()
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        add_styled_paragraph(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = _re.sub(r'[^\w\-]', '_', requirement.title)[:60]
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="spec_{safe_name}.docx"'},
    )


@router.get("/{requirement_id}/history", response_model=list[VersionOut])
def requirement_history(
    requirement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> list[RequirementVersion]:
    del current_user
    _get_requirement(db, requirement_id)
    return (
        db.query(RequirementVersion)
        .filter(RequirementVersion.requirement_id == requirement_id)
        .order_by(RequirementVersion.version_number.desc())
        .all()
    )
